#!/usr/bin/env python3
"""Generate MaintAI Architecture Word document — based on actual codebase analysis."""

from docx import Document
from docx.shared import Inches, Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

# ── colour palette ──────────────────────────────────────────────────────────
DARK_BLUE   = RGBColor(0x0D, 0x1B, 0x3E)
MID_BLUE    = RGBColor(0x1E, 0x40, 0x8F)
ACCENT_CYAN = RGBColor(0x00, 0xB4, 0xD8)
LIGHT_GRAY  = RGBColor(0xF2, 0xF4, 0xF8)
WHITE       = RGBColor(0xFF, 0xFF, 0xFF)
TEXT_DARK   = RGBColor(0x1A, 0x1A, 0x2E)
GREEN_OK    = RGBColor(0x15, 0x80, 0x3D)
RED_ALERT   = RGBColor(0xDC, 0x26, 0x26)
ORANGE      = RGBColor(0xEA, 0x58, 0x0C)
PURPLE      = RGBColor(0x6D, 0x28, 0xD9)

# ── helpers ──────────────────────────────────────────────────────────────────

def set_cell_bg(cell, rgb: RGBColor):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    hex_color = f"{rgb[0]:02X}{rgb[1]:02X}{rgb[2]:02X}"
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), hex_color)
    tcPr.append(shd)


def cell_text(cell, text, bold=False, color=None, size=9, align=WD_ALIGN_PARAGRAPH.LEFT):
    para = cell.paragraphs[0]
    para.alignment = align
    run = para.add_run(text)
    run.bold = bold
    run.font.size = Pt(size)
    if color:
        run.font.color.rgb = color


def heading(doc, text, level=1):
    para = doc.add_paragraph()
    para.paragraph_format.space_before = Pt(18 if level == 1 else 10)
    para.paragraph_format.space_after = Pt(5)
    run = para.add_run(text)
    run.bold = True
    if level == 1:
        run.font.size = Pt(15)
        run.font.color.rgb = MID_BLUE
        run.font.underline = True
    elif level == 2:
        run.font.size = Pt(12)
        run.font.color.rgb = DARK_BLUE
    else:
        run.font.size = Pt(10.5)
        run.font.color.rgb = MID_BLUE
    return para


def body(doc, text, bold=False, color=None, indent=False):
    para = doc.add_paragraph()
    para.paragraph_format.space_after = Pt(4)
    if indent:
        para.paragraph_format.left_indent = Inches(0.25)
    run = para.add_run(text)
    run.font.size = Pt(10)
    run.bold = bold
    if color:
        run.font.color.rgb = color
    return para


def bullet(doc, text, level=0, color=None):
    para = doc.add_paragraph(style='List Bullet')
    para.paragraph_format.left_indent = Inches(0.25 * (level + 1))
    para.paragraph_format.space_after = Pt(2)
    run = para.add_run(text)
    run.font.size = Pt(10)
    if color:
        run.font.color.rgb = color
    return para


def code_block(doc, text):
    para = doc.add_paragraph()
    para.paragraph_format.left_indent = Inches(0.3)
    para.paragraph_format.space_before = Pt(1)
    para.paragraph_format.space_after = Pt(1)
    run = para.add_run(text)
    run.font.name = 'Courier New'
    run.font.size = Pt(8.5)
    run.font.color.rgb = RGBColor(0x2D, 0x6A, 0x4F)
    return para


def divider(doc):
    para = doc.add_paragraph()
    para.paragraph_format.space_before = Pt(2)
    para.paragraph_format.space_after = Pt(2)
    pPr = para._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single')
    bottom.set(qn('w:sz'), '6')
    bottom.set(qn('w:space'), '1')
    bottom.set(qn('w:color'), '1E408F')
    pBdr.append(bottom)
    pPr.append(pBdr)


def add_header_banner(doc, title, subtitle):
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    cell = table.cell(0, 0)
    set_cell_bg(cell, DARK_BLUE)
    para = cell.paragraphs[0]
    para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    para.paragraph_format.space_before = Pt(16)
    para.paragraph_format.space_after = Pt(6)
    r1 = para.add_run(title + '\n')
    r1.bold = True
    r1.font.size = Pt(20)
    r1.font.color.rgb = WHITE
    r2 = para.add_run(subtitle)
    r2.font.size = Pt(10.5)
    r2.font.color.rgb = ACCENT_CYAN
    doc.add_paragraph()


def add_info_box(doc, label, value, label_color=None):
    table = doc.add_table(rows=1, cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    lc = table.cell(0, 0)
    vc = table.cell(0, 1)
    lc.width = Inches(1.9)
    vc.width = Inches(4.6)
    set_cell_bg(lc, label_color or MID_BLUE)
    set_cell_bg(vc, LIGHT_GRAY)
    cell_text(lc, label, bold=True, color=WHITE, size=9)
    cell_text(vc, value, size=9)
    doc.add_paragraph()


def styled_table(doc, headers, rows, col_widths=None):
    n_cols = len(headers)
    table = doc.add_table(rows=1 + len(rows), cols=n_cols)
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table.style = 'Table Grid'
    hdr_row = table.rows[0]
    for i, h in enumerate(headers):
        cell = hdr_row.cells[i]
        set_cell_bg(cell, DARK_BLUE)
        cell_text(cell, h, bold=True, color=WHITE, size=9, align=WD_ALIGN_PARAGRAPH.CENTER)
        if col_widths:
            cell.width = Inches(col_widths[i])
    for ri, row_data in enumerate(rows):
        bg = LIGHT_GRAY if ri % 2 == 0 else WHITE
        for ci, val in enumerate(row_data):
            cell = table.rows[ri + 1].cells[ci]
            set_cell_bg(cell, bg)
            bold = ci == 0
            cell_text(cell, str(val), bold=bold, size=9)
            if col_widths:
                cell.width = Inches(col_widths[ci])
    doc.add_paragraph()
    return table


def ascii_diagram(doc, lines, title=None):
    if title:
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(6)
        r = p.add_run(title)
        r.bold = True
        r.font.size = Pt(9.5)
        r.font.color.rgb = MID_BLUE
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    cell = table.cell(0, 0)
    set_cell_bg(cell, RGBColor(0xF0, 0xF4, 0xFF))
    para = cell.paragraphs[0]
    para.paragraph_format.space_before = Pt(6)
    para.paragraph_format.space_after = Pt(6)
    run = para.add_run('\n'.join(lines))
    run.font.name = 'Courier New'
    run.font.size = Pt(7.5)
    run.font.color.rgb = DARK_BLUE
    doc.add_paragraph()


def callout(doc, text, color=ORANGE):
    table = doc.add_table(rows=1, cols=1)
    cell = table.cell(0, 0)
    set_cell_bg(cell, RGBColor(0xFF, 0xF8, 0xEC))
    para = cell.paragraphs[0]
    para.paragraph_format.space_before = Pt(4)
    para.paragraph_format.space_after = Pt(4)
    para.paragraph_format.left_indent = Inches(0.1)
    run = para.add_run(text)
    run.font.size = Pt(9)
    run.font.color.rgb = color
    doc.add_paragraph()


# ════════════════════════════════════════════════════════════════════════════
# DOCUMENT START
# ════════════════════════════════════════════════════════════════════════════
doc = Document()
section = doc.sections[0]
section.left_margin   = Inches(1.0)
section.right_margin  = Inches(1.0)
section.top_margin    = Inches(0.8)
section.bottom_margin = Inches(0.8)

# ── COVER ────────────────────────────────────────────────────────────────────
add_header_banner(
    doc,
    "MaintAI — Documento di Architettura Tecnica",
    "Versione 3.3.1  |  Giugno 2026  |  Uso interno IT / DevOps"
)

for label, value in [
    ("Prodotto",        "MaintAI — Sistema di Gestione Manutenzione Industriale AI-Powered"),
    ("Versione",        "3.3.1 (build 2026-06-02)"),
    ("Audience",        "Team IT, DevOps, Security Officer, System Architect"),
    ("Classificazione", "Interno — non distribuire"),
    ("Data documento",  "06 Giugno 2026"),
]:
    add_info_box(doc, label, value)

doc.add_paragraph()
divider(doc)

# ── INDICE ───────────────────────────────────────────────────────────────────
heading(doc, "INDICE", 2)
for item in [
    "1.  Panoramica del sistema",
    "2.  Stack tecnologico completo",
    "3.  Architettura applicativa — diagramma a livelli",
    "4.  Frontend — Next.js 16 + React 19",
    "5.  Backend — FastAPI (33 router)",
    "6.  Services Layer — 22 servizi",
    "7.  Database e persistenza",
    "8.  Migrazioni Alembic (16 versioni)",
    "9.  Multi-tenancy — architettura e isolamento",
    "10. Autenticazione e autorizzazione (JWT + Cookie)",
    "11. Sicurezza — OWASP Top 10 e misure adottate",
    "12. Piano AI Felix — motore di pianificazione",
    "13. FIE — Failure Identification Engine (FMECA)",
    "14. Integrazioni esterne",
    "15. Modelli dati ORM (24 modelli)",
    "16. Deploy e infrastruttura (Vercel + Render + Tauri)",
    "17. Variabili d'ambiente",
    "18. Logging e osservabilità",
]:
    bullet(doc, item)

doc.add_page_break()

# ════════════════════════════════════════════════════════════════════════════
# 1. PANORAMICA
# ════════════════════════════════════════════════════════════════════════════
heading(doc, "1. Panoramica del Sistema")
body(doc,
     "MaintAI è una piattaforma SaaS multi-tenant per la gestione della manutenzione industriale "
     "in impianti manifatturieri, energetici e portuali. Combina un motore di pianificazione "
     "deterministico con un layer AI (OpenAI GPT-4) per ottimizzare la schedulazione degli "
     "interventi, la diagnosi guasti (FIE/FMECA), la manutenzione su condizione e l'estrazione "
     "automatica di piani da manuali PDF. È disponibile come web app, app desktop nativa (Tauri 2) "
     "e vista mobile.")

body(doc, "Utenti principali:", bold=True)
bullet(doc, "Responsabile Manutenzione / Planner — KPI, pianificazione AI, coordinamento globale")
bullet(doc, "Tecnico sul campo — piano lavori, esecuzione ticket, diagnosi AI, firma digitale")
bullet(doc, "SuperAdmin — gestione multi-tenant, moduli, configurazione di sistema")

body(doc, "Funzionalità core (verificate nel codice):", bold=True)
for feat in [
    "Gerarchia: Siti → Impianti → Asset con dati tecnici, QR code, criticità A/B/C",
    "Ticket con 5 stati, tipo CM/PM/BD, priorità, allegati, firma digitale, export Excel",
    "Piano AI Felix — motore deterministico + GPT, viste Gantt/Kanban/Calendario, rolling planner",
    "FIE — Failure Identification Engine: analisi FMECA, knowledge base guasti, learning loop",
    "Manutenzione su condizione: trigger basati su metriche (ore funzionamento, soglie)",
    "Sessione diagnostica AI interattiva (RCA via OpenAI, history conversazione)",
    "Caricamento manuali PDF → estrazione automatica piano manutenzione (GPT-4 Vision)",
    "Dashboard KPI real-time (polling 30s), grafici Recharts",
    "Kanban board drag-and-drop (@dnd-kit), Gantt risorse, mappa tecnici (Leaflet)",
    "Email-to-Ticket via IMAP polling ogni 5 minuti (backoff esponenziale su errori)",
    "Check Primo Livello — checklist ispezione pubblica via token UUID (no auth)",
    "Attestati e compliance tecnici, bulk import da Excel, report economico/ROI",
    "Desktop app Tauri 2 (Rust wrapper, WebView2/WebKit) — Windows/macOS/Linux",
    "Auto-generazione ticket da piani ricorrenti (AttivitaManutenzione)",
    "Retention job: soft-delete automatico ticket/piani dopo N giorni",
]:
    bullet(doc, feat)

divider(doc)

# ════════════════════════════════════════════════════════════════════════════
# 2. STACK TECNOLOGICO
# ════════════════════════════════════════════════════════════════════════════
heading(doc, "2. Stack Tecnologico Completo")

styled_table(doc,
    headers=["Livello", "Tecnologia", "Versione", "Note"],
    rows=[
        ("Frontend",          "Next.js (App Router)",      "16.1.6",          "React 19.2.3, TypeScript 5"),
        ("UI / Stili",        "Tailwind CSS v4 + shadcn/ui","4.2.1",          "Dark industrial design system"),
        ("Grafici",           "Recharts",                  "3.8.1",           "Dashboard KPI real-time"),
        ("Drag & Drop",       "@dnd-kit/core + sortable",  "6.3.1 / 10.0.0", "Kanban board planning"),
        ("Tabelle",           "@tanstack/react-table",     "8.21.3",          "Tabelle dati con sort/filter"),
        ("Mappe",             "Leaflet + react-leaflet",   "1.9.4 / 5.0.0",  "Mappa tecnici in campo"),
        ("Toast / Notify",    "Sonner",                    "2.0.7",           "Notifiche UI"),
        ("Date utility",      "date-fns",                  "4.1.0",           "Gestione date/orari"),
        ("Desktop wrapper",   "Tauri 2",                   "2.x (Rust)",      "App nativa Win/Mac/Linux"),
        ("Backend",           "FastAPI",                   "0.135.1",         "Python 3.12, ASGI"),
        ("ASGI Server",       "Uvicorn [standard]",        "0.42.0",          "Produzione via Render"),
        ("ORM",               "SQLAlchemy",                "2.0.48",          "Sync sessions, multi-DB"),
        ("Migrations",        "Alembic",                   "1.18.4",          "16 versioni, batch SQLite"),
        ("Validazione",       "Pydantic v2",               "2.12.5",          "Schema request/response"),
        ("Rate Limiting",     "SlowAPI",                   "0.1.9",           "Wrapper su Starlette"),
        ("JWT",               "PyJWT",                     "2.12.1",          "HS256, blacklist jti"),
        ("Password hash",     "bcrypt / passlib",          "5.0.0 / 1.7.4",  "Salt random per utente"),
        ("Crittografia",      "cryptography (Fernet)",     "46.0.7",          "Cifra password IMAP in DB"),
        ("AI / LLM",          "OpenAI API",                ">=1.0.0 <2.0.0", "GPT-4.1 planner, diagnosi, PDF"),
        ("PDF",               "PyPDF2 + pdf2image",        "3.0.1 / 1.17+",  "Estrazione testo + immagini"),
        ("QR Code",           "Segno",                     "1.6.6",           "Generazione QR asset (PNG/SVG)"),
        ("Email",             "imap-tools",                "1.11.1",          "IMAP polling ticket"),
        ("Excel",             "openpyxl",                  "3.1.5",           "Import bulk + export"),
        ("HTTP client",       "httpx",                     "0.28.1",          "Chiamate API async"),
        ("DB locale",         "SQLite 3",                  "3.x",             "Dev + ambiente demo"),
        ("DB cloud",          "PostgreSQL",                "15.x",            "Render managed"),
        ("Storage file",      "Supabase Storage",          ">=2.0.0",         "Allegati ticket e manuali"),
        ("Deploy frontend",   "Vercel",                    "—",               "Edge CDN, deploy auto da git"),
        ("Deploy backend",    "Render",                    "—",               "Container web service"),
        ("Meteo",             "Open-Meteo API",            "—",               "Vincoli meteo asset outdoor"),
        ("Testing",           "pytest + httpx",            "9.0.3",           "Unit + integration test"),
    ],
    col_widths=[1.5, 1.8, 1.1, 2.1]
)

divider(doc)

# ════════════════════════════════════════════════════════════════════════════
# 3. ARCHITETTURA A LIVELLI
# ════════════════════════════════════════════════════════════════════════════
heading(doc, "3. Architettura Applicativa — Diagramma a Livelli")

ascii_diagram(doc, [
    "┌─────────────────────────────────────────────────────────────────────────────┐",
    "│              CLIENT LAYER  (Browser / Desktop Tauri 2)                      │",
    "│                                                                             │",
    "│  ┌─────────────────────────────┐    ┌───────────────────────────────────┐  │",
    "│  │   WEB APP (Vercel CDN)      │    │   DESKTOP APP (Tauri 2 / Rust)    │  │",
    "│  │   Next.js 16 App Router     │    │   WebView2/WebKit + Next.js       │  │",
    "│  │   React 19 + TypeScript     │    │   Windows / macOS / Linux         │  │",
    "│  └───────────────┬─────────────┘    └────────────────┬──────────────────┘  │",
    "│                  │                                   │                     │",
    "│     AuthContext  │  JWT Cookie (HttpOnly, SameSite)  │  stesso frontend    │",
    "│     api.ts       │  timeout: 30s std / 120s AI       │                     │",
    "└──────────────────┼───────────────────────────────────┼─────────────────────┘",
    "                   │         HTTPS / REST              │",
    "┌──────────────────▼───────────────────────────────────▼─────────────────────┐",
    "│                    BACKEND API LAYER  (Render — Python 3.12)                │",
    "│                    FastAPI 0.135  +  Uvicorn 0.42  (ASGI)                  │",
    "│                                                                             │",
    "│  ┌────────────────────────────────────────────────────────────────────┐    │",
    "│  │  Middleware Stack                                                   │    │",
    "│  │  CORS (whitelist) → Anti-CSRF → Security Headers → Rate Limit      │    │",
    "│  └──────────────────────────────┬─────────────────────────────────────┘    │",
    "│                                 │                                          │",
    "│  ┌──────────────────────────────▼─────────────────────────────────────┐    │",
    "│  │  33 Router Modules  (api/routes/*)           + 14 V1 mirrors       │    │",
    "│  │  auth │ tickets │ planning │ assets │ tecnici │ diagnostic │ ...   │    │",
    "│  │  Module-gated: is_module_enabled() per feature toggles             │    │",
    "│  └──────────────────────────────┬─────────────────────────────────────┘    │",
    "│                                 │                                          │",
    "│  ┌──────────────────────────────▼─────────────────────────────────────┐    │",
    "│  │  Services Layer  (22 servizi)                                       │    │",
    "│  │  planner_engine │ ai_planner │ failure_engine │ pdf_service │ ...  │    │",
    "│  │  email_poller │ auto_ticket │ retention │ condition_maintenance    │    │",
    "│  └──────────────────────────────┬─────────────────────────────────────┘    │",
    "│                                 │                                          │",
    "│  ┌──────────────────────────────▼─────────────────────────────────────┐    │",
    "│  │  ORM Layer  —  SQLAlchemy 2.0  (24 modelli ORM)                    │    │",
    "│  │  Repositories pattern per alcuni domini                            │    │",
    "│  │  Pydantic v2 schemas per request/response validation               │    │",
    "│  └──────────────┬─────────────────────────┬───────────────────────────┘    │",
    "└─────────────────┼─────────────────────────┼────────────────────────────────┘",
    "                  │                         │",
    "  ┌───────────────▼──────────┐   ┌──────────▼────────────────────────────┐",
    "  │  PostgreSQL 15 (Render)  │   │  API Esterne                          │",
    "  │  — Produzione            │   │  OpenAI GPT-4.1 (planner, diagnosi)  │",
    "  │  — Backup automatico     │   │  OpenAI GPT-4.1-mini (PDF, guide)    │",
    "  │  — SSL enforced          │   │  Open-Meteo (meteo asset outdoor)    │",
    "  │                          │   │  IMAP server (email → ticket)        │",
    "  │  SQLite (dev + demo)     │   │  Supabase (storage allegati)         │",
    "  └──────────────────────────┘   └───────────────────────────────────────┘",
], title="Figura 1 — Architettura a livelli MaintAI v3.3.1")

divider(doc)

# ════════════════════════════════════════════════════════════════════════════
# 4. FRONTEND
# ════════════════════════════════════════════════════════════════════════════
heading(doc, "4. Frontend — Next.js 16 + React 19")

body(doc, "Il frontend usa Next.js 16 App Router con rendering ibrido (SSR + CSR). "
     "Il JWT è gestito via cookie HttpOnly (SameSite=Lax) — non in localStorage — per protezione XSS.")

heading(doc, "4.1 Pagine principali (27 route)", 3)
styled_table(doc,
    headers=["Route URL", "File", "Scopo"],
    rows=[
        ("/",                    "page.tsx",                           "Homepage / redirect dashboard"),
        ("/login",               "login/page.tsx",                    "Autenticazione — emit JWT cookie"),
        ("/dashboard",           "dashboard/page.tsx",                "KPI real-time, grafici Recharts"),
        ("/ticket",              "ticket/page.tsx",                   "Lista ticket paginata, filtri, Kanban"),
        ("/assets",              "assets/page.tsx",                   "Inventario asset, ricerca QR/codice"),
        ("/impianti",            "impianti/page.tsx",                 "Gestione impianti (plant/facility)"),
        ("/tecnici",             "tecnici/page.tsx",                  "Anagrafica tecnici + assenze"),
        ("/planning",            "planning/page.tsx",                 "AI Planning: Gantt, Kanban, Rolling"),
        ("/planning/risorse",    "planning/risorse/page.tsx",         "Allocazione risorse tecnici"),
        ("/piani",               "piani/page.tsx",                    "Piani manutenzione generati"),
        ("/piani-manutenzione",  "piani-manutenzione/page.tsx",       "Creazione piani strutturati multi-asset"),
        ("/manuali",             "manuali/page.tsx",                  "Upload PDF, estrazione piano AI"),
        ("/scadenze",            "scadenze/page.tsx",                 "Scadenze ticket e attività"),
        ("/condizioni",          "condizioni/page.tsx",               "Manutenzione su condizione"),
        ("/diagnostic",          "diagnostic/page.tsx",               "FIE: diagnosi AI FMECA"),
        ("/compliance",          "compliance/page.tsx",               "Compliance, audit trail"),
        ("/storico/[asset_id]",  "storico/[asset_id]/page.tsx",       "Storico interventi per asset"),
        ("/report/economico",    "report/economico/page.tsx",         "Report costi / ROI"),
        ("/admin/utenti",        "admin/utenti/page.tsx",             "Gestione utenti (superadmin)"),
        ("/admin/tenants",       "admin/tenants/page.tsx",            "Gestione tenant (superadmin)"),
        ("/admin/bulk-import",   "admin/bulk-import/page.tsx",        "Import bulk da Excel"),
        ("/admin/email",         "admin/email/page.tsx",              "Config IMAP per tenant"),
        ("/admin/logs",          "admin/logs/page.tsx",               "SystemLog (superadmin)"),
        ("/profilo",             "profilo/page.tsx",                  "Profilo utente"),
        ("/mobile",              "mobile/page.tsx",                   "Vista mobile ottimizzata"),
        ("/check/[token]",       "check/[token]/page.tsx",            "Check primo livello (pubblico, no auth)"),
        ("/print/asset-qr/[id]", "print/asset-qr/[id]/page.tsx",     "Stampa QR code asset"),
    ],
    col_widths=[2.1, 2.2, 2.2]
)

heading(doc, "4.2 Componenti principali", 3)
styled_table(doc,
    headers=["Componente", "Funzione"],
    rows=[
        ("KanbanBoard.tsx",         "Drag-and-drop ticket (@dnd-kit) con colonne stato"),
        ("GanttGiornaliero.tsx",    "CSS grid 18 slot 08:00–17:00 per tecnico"),
        ("GanttRisorse.tsx",        "Vista Gantt per allocazione risorse"),
        ("CalendarioMensile.tsx",   "Griglia mensile con dots colorati per tipo ticket"),
        ("BadgeEfficienza.tsx",     "Conic-gradient + breakdown 5 componenti efficiency score"),
        ("EmergencyMap.tsx",        "Mappa Leaflet posizione tecnici in campo"),
        ("QrScanner.tsx",           "Scanner QR code (mobile) per identificazione asset"),
        ("SignaturePad.tsx",        "Firma digitale tecnico per chiusura ticket"),
        ("VoiceRecorder.tsx",       "Registrazione audio memo per ticket"),
        ("WeatherWidget.tsx",       "Meteo real-time Open-Meteo per asset outdoor"),
        ("GuideBot.tsx",            "Chat AI assistente (streaming OpenAI)"),
        ("NotificationPanel.tsx",   "Notifiche in-app real-time"),
        ("GlobalQuickTicket.tsx",   "Creazione rapida ticket da qualsiasi pagina"),
        ("BackendStatus.tsx",       "Health check backend con indicatore visivo"),
        ("StatusToggle.tsx",        "Toggle stato asset service ↔ fermo_prog inline"),
        ("UploadAllegati.tsx",      "Upload file allegati per ticket e asset"),
    ],
    col_widths=[2.5, 4.0]
)

heading(doc, "4.3 Comunicazione con il backend", 3)
bullet(doc, "JWT trasportato in cookie HttpOnly (SameSite=Lax, Secure in prod) — non esponibile via JS")
bullet(doc, "Client fetch centralizzato in lib/api.ts con timeout adattivo: 30s standard, 120s endpoint AI")
bullet(doc, "Anti-CSRF: middleware backend verifica Origin/Referer obbligatori per POST/PUT/DELETE")
bullet(doc, "NEXT_PUBLIC_API_BASE configura la base URL via variabile Vercel")
bullet(doc, "Versione frontend: 3.3.0 (build 2026-05-25) — sfasata di 1 patch vs backend 3.3.1")

callout(doc, "⚠  NOTA: La versione in frontend/app/lib/version.ts (3.3.0) è sfasata di 1 patch "
        "rispetto al backend (3.3.1 / 2026-06-02). Allineare prima del prossimo deploy.")

divider(doc)

# ════════════════════════════════════════════════════════════════════════════
# 5. BACKEND — ROUTER
# ════════════════════════════════════════════════════════════════════════════
heading(doc, "5. Backend — FastAPI (33 Router + 14 Mirror V1)")

body(doc,
     "Il backend FastAPI registra 33 router legacy (senza prefisso /v1) e 14 router mirror con prefisso /v1 "
     "per migrazione graduale del frontend. Tutti (eccetto health, auth, modules e check_primo_livello) "
     "richiedono JWT valido e sono controllati da is_module_enabled() per feature toggle per tenant.")

styled_table(doc,
    headers=["Router", "URL prefix", "Modulo funzionale", "Note chiave"],
    rows=[
        ("health",              "/health",              "Core",                 "Pubblico — versione, DB status, uptime"),
        ("auth",                "/auth",                "Core",                 "Login, logout, register, refresh"),
        ("modules",             "/modules",             "Core",                 "Feature toggle per tenant"),
        ("dashboard",           "/dashboard",           "Dashboard",            "KPI aggregati, metriche real-time"),
        ("assets",              "/assets",              "Asset Management",     "CRUD asset, QR, criticità, meteo"),
        ("impianti",            "/impianti",            "Asset Management",     "CRUD impianti (plant/facility)"),
        ("siti",                "/siti",                "Asset Management",     "CRUD siti geografici"),
        ("procedure",           "/procedure",           "Asset Management",     "Procedure operative per asset (passi JSON)"),
        ("note_asset",          "/note_asset",          "Asset Management",     "Note tecniche senior (upsert)"),
        ("check_primo_livello", "/check_primo_livello", "Asset Management",     "Checklist pubblica via UUID token"),
        ("asset_documenti",     "/asset_documenti",     "Asset Management",     "Documenti/esplosi allegati asset (BYTEA)"),
        ("tecnici",             "/tecnici",             "Technicians",          "CRUD tecnici, competenze, assenze, orari"),
        ("tickets",             "/tickets",             "Tickets",              "CRUD ticket, stati, allegati, export Excel"),
        ("scadenze",            "/scadenze",            "Deadlines",            "Scadenze ticket e attività ricorrenti"),
        ("logs",                "/logs",                "System Logs",          "Lettura SystemLog (admin)"),
        ("scheduler",           "/scheduler",           "Planning",             "Scheduler legacy (redirect a planning)"),
        ("planning",            "/planning",            "Planning",             "Generate, conferma, storico, deautorizza"),
        ("manuali",             "/manuali",             "Manuals",              "Upload PDF, parsing, estrazione piano"),
        ("diagnostic",          "/diagnostic",          "Diagnostic AI",        "Sessioni RCA interattive (OpenAI stream)"),
        ("problem_analysis",    "/problem_analysis",    "Diagnostic AI",        "Analisi problema multi-step"),
        ("failure_engine",      "/failure_engine",      "Diagnostic AI",        "FIE FMECA: knowledge base + scoring"),
        ("piani",               "/piani",               "Maintenance Plans",    "Piani legacy (compatibilità)"),
        ("piano_manutenzione",  "/piano_manutenzione",  "Maintenance Plans",    "Piani strutturati multi-asset"),
        ("tenants",             "/tenants",             "Tenant Admin",         "Gestione clienti (solo superadmin)"),
        ("email_config",        "/email_config",        "Email-to-Ticket",      "Config IMAP per tenant"),
        ("bulk_import",         "/bulk_import",         "Bulk Import",          "Import massal da Excel (openpyxl)"),
        ("utenti",              "/utenti",              "User Admin",           "Gestione utenti del tenant"),
        ("desktop_update",      "/desktop_update",      "Desktop Updates",      "Auto-update app Tauri (delta update)"),
        ("conditions",          "/conditions",          "Condition Maintenance", "Letture metriche, trigger condition-based"),
        ("guide",               "/guide",               "Guide AI",             "Chat AI assistente (GuideBot streaming)"),
        ("attestati",           "/attestati",           "Compliance",           "Qualifiche e certificazioni tecnici"),
        ("report",              "/report",              "Economic Reports",     "Report costi, ROI, analisi economica"),
        ("emergency",           "/emergency",           "Emergency",            "Gestione emergenze, mappa tecnici"),
    ],
    col_widths=[1.5, 1.5, 1.5, 2.0]
)

heading(doc, "5.1 Startup sequence (lifespan)", 3)
body(doc, "All'avvio del backend FastAPI, il lifespan context manager esegue in sequenza:")
for step in [
    "1. Setup logging (Python standard + file handler)",
    "2. init_backend() — carica config, verifica env vars obbligatorie",
    "3. Check DB connection (retry con backoff esponenziale, max 5 tentativi)",
    "4. alembic upgrade head — applica migration mancanti",
    "5. init_db() — seed dati iniziali se DB vuoto",
    "6. _ensure_columns() — fallback DDL idempotente per cloud deploy",
    "7. Background tasks: retention_job, email_poller_task, auto_ticket_job (se moduli abilitati)",
]:
    bullet(doc, step)

heading(doc, "5.2 Middleware stack", 3)
styled_table(doc,
    headers=["Middleware", "Funzione"],
    rows=[
        ("CORS",            "Origin whitelist: prod Vercel + Tauri + localhost dev"),
        ("Anti-CSRF",       "Verifica Origin/Referer obbligatori per POST/PUT/DELETE — fail-closed"),
        ("Security Headers","X-Content-Type-Options, Referrer-Policy, Permissions-Policy, HSTS (prod)"),
        ("Rate Limiting",   "SlowAPI — limiti configurabili per endpoint (es. 10/min su /auth/login)"),
        ("Exception Handler","AppError custom + handler generico — risposta JSON strutturata sempre"),
    ],
    col_widths=[1.8, 4.7]
)

divider(doc)

# ════════════════════════════════════════════════════════════════════════════
# 6. SERVICES LAYER
# ════════════════════════════════════════════════════════════════════════════
heading(doc, "6. Services Layer — 22 Servizi")

styled_table(doc,
    headers=["Servizio", "File", "Scopo"],
    rows=[
        ("AI Planner (Felix)",      "ai_planner_service.py",         "Pianificazione AI via OpenAI + vincoli meteo Open-Meteo"),
        ("PlannerEngine",           "planner_engine.py",             "Motore deterministico puro (dataclass, no ORM, testabile)"),
        ("PlannerEngine Bridge",    "planner_engine_bridge.py",      "Adapter ORM → PlannerEngine → plan_json"),
        ("Rolling Planner",         "rolling_planner_engine.py",     "Rigenera piani ogni giorno (pianificazione continuativa)"),
        ("Adaptive Estimator",      "adaptive_estimator.py",         "Corregge durata ticket basandosi su storico PlannerFeedback"),
        ("Failure Engine (FIE)",    "failure_engine.py",             "FMECA: analisi sintomi vs knowledge base, scoring RPN"),
        ("Condition Maintenance",   "condition_maintenance_service.py","Trigger manutenzione su soglie metriche asset"),
        ("Opportunistic Service",   "opportunistic_service.py",      "Suggerisce task correlati quando asset è già in manutenzione"),
        ("Auto Ticket Service",     "auto_ticket_service.py",        "Genera automaticamente ticket da AttivitaManutenzione"),
        ("Retention Service",       "retention_service.py",          "Soft-delete automatico ticket/piani dopo N giorni"),
        ("Email Poller",            "email_poller.py",               "IMAP polling ogni 5 min → crea Ticket (backoff su errori)"),
        ("PDF Service",             "pdf_service.py",                "PDF → testo + immagini (PyPDF2 + pdf2image)"),
        ("QR Service",              "qr_service.py",                 "Generazione QR code PNG base64 per asset"),
        ("Weather Service",         "weather_service.py",            "Fetch previsioni meteo Open-Meteo per vincoli asset"),
        ("Ticket Service",          "ticket_service.py",             "CRUD ticket + business logic (stati, assegnazioni)"),
        ("Scheduler Service",       "scheduler_service.py",          "Coordinamento scheduler background jobs"),
        ("Assets Service",          "assets_service.py",             "CRUD asset con repository pattern"),
        ("OpenAI Service",          "ai/openai_service.py",          "Wrapper OpenAI — modello configurabile, caching"),
        ("Anonymization Service",   "ai/anonymization_service.py",   "Anonimizza dati sensibili prima invio a OpenAI"),
        ("Diagnostic Service",      "ai/diagnostic_service.py",      "Diagnosi AI: analisi sintomi + storia conversazione"),
        ("Manuals AI Service",      "ai/manuals_ai_service.py",      "Estrazione concetti manuali PDF via GPT-4 Vision"),
        ("Problem Analysis Service","ai/problem_analysis_service.py","Analisi problema multi-step con AI"),
    ],
    col_widths=[2.0, 2.1, 2.4]
)

divider(doc)

# ════════════════════════════════════════════════════════════════════════════
# 7. DATABASE
# ════════════════════════════════════════════════════════════════════════════
heading(doc, "7. Database e Persistenza")

styled_table(doc,
    headers=["Ambiente", "Database", "URL / Path", "Note"],
    rows=[
        ("Produzione cloud", "PostgreSQL 15", "Render managed DB (SSL enforced)", "Multi-tenant, backup daily automatico"),
        ("Demo / Sandbox",   "SQLite 3",      "demo.db (locale)",                 "Utenti demo isolati, reset rapido"),
        ("Sviluppo locale",  "SQLite 3",      "maintai.db",                       "Nessuna infrastruttura necessaria"),
    ],
    col_widths=[1.6, 1.5, 2.2, 1.7]
)

heading(doc, "7.1 Routing automatico del DB", 3)
body(doc,
     "La funzione get_db() in backend/core/dependencies.py usa SessionLocal (connesso a PostgreSQL o SQLite "
     "in base all'env DATABASE_URL). Il routing demo/prod non avviene più nel get_db(), ma è gestito "
     "dalla configurazione dell'engine: utenti con is_demo=True nel JWT usano DEMO_DATABASE_URL.")

heading(doc, "7.2 Pattern di isolamento tenant nelle query", 3)
code_block(doc, "# Ogni query filtra obbligatoriamente per tenant_id:")
code_block(doc, "tenant_id = get_current_tenant_id(payload, x_tenant_id)")
code_block(doc, "tickets = db.query(Ticket).filter(Ticket.tenant_id == tenant_id).all()")
code_block(doc, "")
code_block(doc, "# Per risorse singole si usa check_tenant_ownership():")
code_block(doc, "ticket = check_tenant_ownership(db, Ticket, ticket_id, tenant_id)  # 404 se non appartiene")

divider(doc)

# ════════════════════════════════════════════════════════════════════════════
# 8. MIGRAZIONI ALEMBIC
# ════════════════════════════════════════════════════════════════════════════
heading(doc, "8. Migrazioni Alembic — 16 Versioni")

styled_table(doc,
    headers=["#", "Migration ID / File", "Contenuto"],
    rows=[
        ("1",  "fa1c5b675293_initial_master_baseline",          "Schema base iniziale — tutte le tabelle core"),
        ("2",  "35093982921a_add_system_logs_and_email_config", "SystemLog, EmailConfig tables"),
        ("3",  "92deed855516_ristrutturazione_piani_multiasset","Piani multi-asset — tabella associazione M2M"),
        ("4",  "3b5b37e1e43f_m1_m2_asset_criticita_costo_fermo","Campi criticità A/B/C, costo_orario_fermo asset"),
        ("5",  "20260405_add_planning_fields",                  "Asset: weather_constraint, fermo_on_schedule, lat/lon"),
        ("6",  "20260405002_add_ticket_splitting_fields",       "Ticket splitting: parent_ticket_id, is_continuation"),
        ("7",  "20260405003_add_plan_storico_fields",           "Storico piani: plan_number, confirmed_by, deauthorized_at"),
        ("8",  "20260406001_add_scadenza_and_assenza_tenant",   "Scadenza su generated_plans, tenant isolation assenze"),
        ("9",  "20260406002_add_composite_tenant_indexes",      "Indici compositi per performance multi-tenant"),
        ("10", "20260406003_add_audit_trail_and_indexes",       "Audit trail: created_by, closed_by, soft-delete deleted_at"),
        ("11", "20260406004_add_soft_deletion_and_multi_tech",  "Soft deletion completo, multi-tecnico (tecnici_richiesti)"),
        ("12", "20260411001_add_task_plan_link_and_v250_fields","Task ↔ Piano link (piano_id), unificazione v2.5.0"),
        ("13", "c8b3fc6efac0_m4_m5_procedure_note_asset_check", "Tabelle M4/M5: Procedura, NotaAsset, CheckPrimoLivello, Attestato"),
        ("14", "fa2e7d5c222b_add_qr_token_expiry_and_active",   "QR token scadenza — token_expires_at, token_active"),
        ("15", "20260601001_tenant_id_not_null_safe",           "Preparazione constraint tenant_id NOT NULL"),
        ("16", "20260601002_tenant_id_not_null_constraint",     "Applicazione definitiva NOT NULL su tenant_id"),
    ],
    col_widths=[0.3, 2.8, 3.4]
)

body(doc, "Note tecniche Alembic:", bold=True)
bullet(doc, "batch_alter_table per compatibilità SQLite (PostgreSQL usa ALTER TABLE standard)")
bullet(doc, "Transazioni isolate per ogni DDL — evita 'aborted transaction state' su PostgreSQL")
bullet(doc, "Ogni migration è idempotente (IF NOT EXISTS + exception handling)")
bullet(doc, "Fallback _ensure_columns() in main.py: SQL diretto idempotente per cloud deploy")

divider(doc)

# ════════════════════════════════════════════════════════════════════════════
# 9. MULTI-TENANCY
# ════════════════════════════════════════════════════════════════════════════
heading(doc, "9. Multi-Tenancy — Architettura e Isolamento")

body(doc,
     "MaintAI usa un'architettura shared-database / shared-schema. Tutte le tabelle "
     "contengono tenant_id (FK → tenants.id NOT NULL). L'isolamento è garantito a livello "
     "applicativo: ogni query filtra per tenant_id estratto dal JWT. "
     "Dal migration #16, tenant_id è NOT NULL su tutte le tabelle.")

ascii_diagram(doc, [
    "  JWT Cookie (HttpOnly)         →   backend/core/security.py",
    "  ┌───────────────────────────┐",
    "  │  sub: 'mario.rossi'       │     get_current_tenant_id(payload, x_tenant_id_header)",
    "  │  tenant_id: 42   ◄── KEY │     └─→ restituisce 42",
    "  │  ruolo: 'responsabile'    │",
    "  │  is_demo: false           │     Ogni router usa Depends(get_current_tenant_id)",
    "  │  jti: 'uuid-del-token'    │     ↓",
    "  │  token_version: 3         │     Query SQLAlchemy sempre filtrata:",
    "  │  exp: 1749081600          │     .filter(Model.tenant_id == 42)",
    "  └───────────────────────────┘",
    "",
    "  SuperAdmin può override via header HTTP:",
    "  X-Tenant-Id: 99   →   get_current_tenant_id() ritorna 99 (impersonazione)",
], title="Figura 2 — Isolamento dati multi-tenant")

styled_table(doc,
    headers=["Ruolo JWT", "Accesso dati", "Permessi"],
    rows=[
        ("superadmin",    "Tutti i tenant (via X-Tenant-Id)", "CRUD completo + gestione tenant/utenti"),
        ("responsabile",  "Solo tenant JWT",                   "CRUD completo sulle risorse del proprio tenant"),
        ("tecnico",       "Solo tenant JWT",                   "Lettura planning; esecuzione ticket assegnati; diagnosi"),
        ("viewer",        "Solo tenant JWT",                   "Sola lettura su tutto"),
    ],
    col_widths=[1.5, 2.0, 3.0]
)

body(doc, "Meccanismi di protezione aggiuntivi:", bold=True)
bullet(doc, "RevokedToken table — blacklist JTI per invalidazione istantanea (logout)")
bullet(doc, "token_version su Utente — invalidazione di tutti i token precedenti (cambio password)")
bullet(doc, "check_tenant_ownership() — verifica ownership prima di ogni modifica a risorsa singola")
bullet(doc, "Indici compositi (tenant_id, id) su tutte le tabelle per performance query")

divider(doc)

# ════════════════════════════════════════════════════════════════════════════
# 10. AUTENTICAZIONE
# ════════════════════════════════════════════════════════════════════════════
heading(doc, "10. Autenticazione e Autorizzazione (JWT + Cookie)")

ascii_diagram(doc, [
    "  Browser                         FastAPI /auth/login",
    "     │                                    │",
    "     │── POST /auth/login ───────────────►│",
    "     │   {username, password}             │",
    "     │                                    │── verifica bcrypt hash (passlib)",
    "     │                                    │── controlla Utente.is_active + Tenant.is_active",
    "     │                                    │── genera JWT (HS256)",
    "     │                                    │   payload: sub, tenant_id, ruolo,",
    "     │                                    │            is_demo, jti, token_version, exp",
    "     │◄── Set-Cookie: maintai_jwt='...'  ─│",
    "     │    HttpOnly=True, SameSite=Lax     │",
    "     │    Secure=True (prod)              │",
    "     │                                    │",
    "     │── GET /tickets ──────────────────►│",
    "     │   Cookie: maintai_jwt='...'        │",
    "     │   (automatico, no JS)              │── decode JWT (PyJWT HS256)",
    "     │                                    │── controlla jti non in RevokedToken",
    "     │                                    │── controlla token_version == Utente.token_version",
    "     │                                    │── estrae tenant_id, ruolo",
    "     │                                    │── query filtrata per tenant_id",
    "     │◄── [{...ticket data...}] ─────────│",
], title="Figura 3 — Flusso autenticazione JWT via Cookie HttpOnly")

heading(doc, "10.1 Funzioni di sicurezza (core/security.py)", 3)
styled_table(doc,
    headers=["Funzione", "Tipo", "Scopo"],
    rows=[
        ("get_current_user_payload()",  "FastAPI Depends", "Decode JWT cookie, verifica blacklist, ritorna payload"),
        ("get_current_tenant_id()",     "FastAPI Depends", "Estrae tenant_id da payload (o header X-Tenant-Id se superadmin)"),
        ("require_superadmin()",        "FastAPI Depends", "Blocca con HTTP 403 se ruolo != superadmin"),
        ("require_roles(*allowed)",     "Depends factory", "Ammette superadmin sempre + ruoli nella whitelist"),
        ("check_tenant_ownership()",    "Utility",         "Verifica che object_id appartenga al tenant — 404 se no"),
        ("encrypt_data() / decrypt_data()", "Utility",     "Fernet AES-128 CBC — cifra/decifra password IMAP in DB"),
        ("get_password_hash()",         "Utility",         "bcrypt con salt random per password utente"),
        ("verify_password()",           "Utility",         "Confronto sicuro con hash bcrypt"),
        ("create_access_token()",       "Utility",         "Genera JWT con jti UUID + exp da config"),
    ],
    col_widths=[2.3, 1.5, 2.7]
)

body(doc, "Costanti JWT:", bold=True)
bullet(doc, "Algoritmo: HS256")
bullet(doc, "Scadenza: 1440 minuti (24 ore, configurabile)")
bullet(doc, "Cookie name: maintai_jwt")
bullet(doc, "Cookie Secure: True in produzione, False in locale")
bullet(doc, "Cookie SameSite: lax (protegge da CSRF cross-site)")

divider(doc)

# ════════════════════════════════════════════════════════════════════════════
# 11. SICUREZZA
# ════════════════════════════════════════════════════════════════════════════
heading(doc, "11. Sicurezza — OWASP Top 10 e Misure Adottate")

body(doc, "Riferimento: docs/SECURITY_GUIDELINES.md + docs/SECURITY_GUIDELINES_MAINTAI.md + "
     "docs/SECURITY_AUDIT_2026-05-30.md")

styled_table(doc,
    headers=["OWASP", "Rischio", "Implementazione in MaintAI"],
    rows=[
        ("A01 — Broken Access Control",
         "Alta",
         "tenant_id NOT NULL su ogni tabella; check_tenant_ownership(); require_superadmin; "
         "indici compositi; X-Tenant-Id solo per superadmin"),
        ("A02 — Cryptographic Failures",
         "Alta",
         "Password hash bcrypt (salt random); JWT HS256 con SECRET_KEY env; "
         "Fernet (AES-128 CBC) per password IMAP; Cookie HttpOnly+Secure"),
        ("A03 — Injection",
         "Alta",
         "ORM SQLAlchemy con bind parameters (zero raw SQL user-input); "
         "Pydantic v2 con Field(min_length, max_length, pattern) su tutti gli input"),
        ("A04 — Insecure Design",
         "Media",
         "Services layer separato; nessuna query cross-tenant; "
         "anonymization_service.py prima di ogni invio a OpenAI"),
        ("A05 — Security Misconfiguration",
         "Media",
         "CORS origin whitelist; Anti-CSRF middleware fail-closed; "
         "Security Headers su ogni risposta; env vars mai in codice sorgente"),
        ("A06 — Vulnerable Components",
         "Media",
         "Dipendenze pinnate in requirements.txt; versioni aggiornate (cryptography 46.0.7, "
         "bcrypt 5.0.0); nessuna dipendenza EOL"),
        ("A07 — Auth Failures",
         "Alta",
         "Rate limit /auth/login (SlowAPI 10/min); JTI blacklist (RevokedToken); "
         "token_version per invalidazione globale; Cookie SameSite=Lax"),
        ("A08 — Software Integrity",
         "Bassa",
         "Deploy da branch main protetto; build Vercel + Render da repo controllato; "
         "auto-update Tauri firmato (desktop_update router)"),
        ("A09 — Logging Failures",
         "Media",
         "SystemLog persistente in DB (ogni operazione critica); "
         "Python logging standard in parallelo; SystemLog visibile da admin UI"),
        ("A10 — SSRF",
         "Media",
         "Open-Meteo e OpenAI chiamati solo da backend con URL fissi; "
         "nessun redirect verso URL fornito da utente; Supabase via SDK ufficiale"),
    ],
    col_widths=[2.0, 0.7, 3.8]
)

heading(doc, "11.1 Misure specifiche aggiuntive", 3)
bullet(doc, "Anti-CSRF middleware: verifica Origin/Referer obbligatori per ogni metodo non-GET — fail-closed")
bullet(doc, "File allegati: su Supabase Storage bucket privato — mai path diretto pubblico")
bullet(doc, "PDF processing: in-memory, non esposto via URL pubblico")
bullet(doc, "QR code pubblico /check/[token]: UUID randomico, scadenza configurabile (token_expires_at), toggle attivo/disattivo")
bullet(doc, "Dati inviati a OpenAI: passano per anonymization_service.py prima dell'invio")
bullet(doc, "Desktop app Tauri: aggiornamenti firmati via endpoint /desktop_update (no MITM possibile)")

divider(doc)

# ════════════════════════════════════════════════════════════════════════════
# 12. PIANO AI FELIX
# ════════════════════════════════════════════════════════════════════════════
heading(doc, "12. Piano AI Felix — Motore di Pianificazione")

styled_table(doc,
    headers=["Parametro mode", "Motore",                      "Velocità",       "Requisiti"],
    rows=[
        ('"deterministic"', "PlannerEngine (Python puro)",     "< 1 secondo",    "Nessuno"),
        ('"ai"',            "OpenAI GPT-4.1",                  "30–120 secondi", "OPENAI_API_KEY"),
        ('"rolling"',       "RollingPlannerEngine (continuo)", "< 1 secondo",    "Nessuno"),
        ('"auto"',          "Deterministico o AI in base a key","—",             "Default"),
    ],
    col_widths=[1.5, 2.2, 1.6, 1.2]
)

heading(doc, "12.1 PlannerEngine deterministico", 3)
body(doc, "Implementato in backend/services/planner_engine.py con sole dataclass Python (zero ORM). "
     "100% testabile in isolamento via pytest.")
bullet(doc, "Priorità scheduling: BD (Breakdown) > PM (Preventiva) > CM (Correttiva) per giornata")
bullet(doc, "Griglia oraria: 08:00–17:00, 18 slot da 30 minuti")
bullet(doc, "Verifica disponibilità tecnico: assenze, orari, slot già occupati")
bullet(doc, "Verifica competenze: Meccanico, Elettricista + PM/CM/BD come competenze implicite su ogni tecnico attivo")
bullet(doc, "Split WO automatico su più giorni se durata supera lo slot disponibile (is_continuation)")
bullet(doc, "Adaptive Estimator: corregge durata stimata basandosi su PlannerFeedback storico")
bullet(doc, "efficiency_score (0–100) con breakdown: copertura_backlog, utilizzo_tecnici, priorità rispettate, ecc.")

heading(doc, "12.2 Rolling Planner", 3)
body(doc,
     "Il RollingPlannerEngine rigenera automaticamente il piano ogni giorno su un orizzonte mobile, "
     "incorporando nuovi ticket e feedback di esecuzione. Non richiede conferma manuale — è progettato "
     "per operazioni continue senza intervento del planner.")

heading(doc, "12.3 Formato plan_json (output comune)", 3)
code_block(doc, '{')
code_block(doc, '  "planned_workorders": [')
code_block(doc, '    {')
code_block(doc, '      "wo_id": 42, "technician_id": 3,')
code_block(doc, '      "planned_date": "2026-04-07", "time_slot": "08:00-10:00",')
code_block(doc, '      "duration_hours": 2.0, "motivation": "...",')
code_block(doc, '      "is_continuation": false, "parent_wo_id": null')
code_block(doc, '    }')
code_block(doc, '  ],')
code_block(doc, '  "deferred_workorders": [{"wo_id": 5, "reason": "tecnico non disponibile"}],')
code_block(doc, '  "fermo_assets": [{"asset_id": 1, "triggered_by_wo_id": 42}],')
code_block(doc, '  "efficiency_score": 78,')
code_block(doc, '  "efficiency_breakdown": {"copertura_backlog": 85, "utilizzo_tecnici": 70, ...}')
code_block(doc, '}')

heading(doc, "12.4 Flusso conferma piano (POST /planning/confirm/{id})", 3)
bullet(doc, "Verifica piano in stato 'draft' e appartiene al tenant")
bullet(doc, "Per ogni wo in planned_workorders: Ticket.stato → 'Pianificato', assegna tecnico, planned_start/finish")
bullet(doc, "Asset con fermo_on_schedule=True → stato = 'Fermo'")
bullet(doc, "GeneratedPlan.status → 'confirmed' + plan_number progressivo (MAX+1 per tenant)")
bullet(doc, "confirmed_by = username JWT corrente; confirmed_at = now()")
bullet(doc, "PlannerFeedback record creato per ogni WO confermato (learning loop)")

divider(doc)

# ════════════════════════════════════════════════════════════════════════════
# 13. FIE — FAILURE IDENTIFICATION ENGINE
# ════════════════════════════════════════════════════════════════════════════
heading(doc, "13. FIE — Failure Identification Engine (FMECA)")

body(doc,
     "Il FIE è il motore di diagnosi guasti AI di MaintAI. Combina una knowledge base FMECA "
     "(Failure Mode, Effects and Criticality Analysis) con scoring RPN e apprendimento dai feedback "
     "dei tecnici per identificare la causa più probabile di un guasto.")

ascii_diagram(doc, [
    "  Tecnico inserisce sintomi nel Ticket",
    "          │",
    "          ▼  POST /failure_engine/analyze",
    "  ┌───────────────────────────────────────────────────────┐",
    "  │  failure_engine.py                                    │",
    "  │                                                       │",
    "  │  1. Carica FailureMode dal DB (global + tenant)       │",
    "  │  2. Filtra per asset_type e componente                │",
    "  │  3. Calcola probability_score per ogni failure mode:  │",
    "  │     - Match keywords sintomi                          │",
    "  │     - RPN = Severity × Occurrence × Detectability    │",
    "  │     - Peso appreso da DiagnosticLearning (feedback)   │",
    "  │  4. Ordina per rpn_weighted DESC                      │",
    "  │  5. Salva in FailureAnalysis                         │",
    "  └───────────────────────────────────────────────────────┘",
    "          │",
    "          ▼  Tecnico conferma o corregge diagnosi",
    "  DiagnosticLearning.save()  →  aggiorna peso_appreso in FailureMode",
], title="Figura 4 — Flusso FIE (Failure Identification Engine)")

styled_table(doc,
    headers=["Modello", "Scopo", "Campi chiave"],
    rows=[
        ("FailureMode",        "Knowledge base FMECA",           "asset_type, component, failure_mode, severity, occurrence, detectability, RPN, peso_appreso"),
        ("FailureAnalysis",    "Risultati analisi per ticket",   "ticket_id, failure_mode_id, probability_score, rpn_weighted, ai_explanation, selected"),
        ("DiagnosticLearning", "Storico conferme (learning)",    "ticket_id, diagnosed_failure_mode_id, real_cause, action_taken, success"),
    ],
    col_widths=[1.8, 2.0, 2.7]
)

divider(doc)

# ════════════════════════════════════════════════════════════════════════════
# 14. INTEGRAZIONI ESTERNE
# ════════════════════════════════════════════════════════════════════════════
heading(doc, "14. Integrazioni Esterne")

styled_table(doc,
    headers=["Servizio", "Protocollo", "Scopo", "Sicurezza"],
    rows=[
        ("OpenAI API",    "HTTPS REST",  "GPT-4.1: pianificazione AI Felix, GuideBot streaming; GPT-4.1-mini: diagnosi FIE, PDF parsing", "API Key in env; dati anonimizzati prima invio"),
        ("Open-Meteo",    "HTTPS REST",  "Previsioni meteo 7gg per vincoli asset outdoor (WeatherWidget, planner)", "API pubblica, no auth; URL fisso nel codice"),
        ("IMAP Server",   "IMAP/TLS",    "Polling email ogni 5 min → crea Ticket automaticamente (backoff su errori)", "Credenziali Fernet cifrate in DB; TLS obbligatorio"),
        ("Supabase",      "HTTPS SDK",   "Storage allegati ticket, manuali PDF, documenti asset", "Bucket privato; signed URL per download; service role key in env"),
        ("PostgreSQL",    "TCP/TLS 5432","DB produzione (Render managed)", "SSL enforced; IP allowlist Render"),
        ("Tauri Update",  "HTTPS",       "Auto-update delta app desktop (Windows/macOS/Linux)", "Firma digitale aggiornamenti; endpoint /desktop_update"),
    ],
    col_widths=[1.3, 1.2, 2.5, 1.5]
)

divider(doc)

# ════════════════════════════════════════════════════════════════════════════
# 15. MODELLI DATI ORM
# ════════════════════════════════════════════════════════════════════════════
heading(doc, "15. Modelli Dati ORM — 24 Classi + 1 Tabella M2M")

body(doc, "Tutti i modelli sono in backend/db/modelli.py (662 righe). "
     "Ogni modello ha tenant_id FK obbligatorio (NOT NULL dal migration #16).")

styled_table(doc,
    headers=["Modello ORM", "Tabella", "Campi chiave"],
    rows=[
        ("Tenant",               "tenants",               "id, nome, slug, is_active, created_at"),
        ("Utente",               "utenti",                "id, username, password_hash, ruolo (superadmin|responsabile|tecnico), is_active, tenant_id, token_version"),
        ("RevokedToken",         "revoked_tokens",        "id, jti, created_at — blacklist logout JWT"),
        ("Sito",                 "siti",                  "id, nome, ubicazione, città, responsabile, tenant_id"),
        ("Impianto",             "impianti",              "id, nome, tipologia, latitude, longitude, sito_id, tenant_id"),
        ("Asset",                "assets",                "id, nome, codice, stato, criticita (A|B|C), weather_constraint, qr_code_b64, costo_orario_fermo, tenant_id"),
        ("Tecnico",              "tecnici",               "id, utente_id, nome, competenze, ore_giornaliere, orario_inizio/fine, sede_indirizzo, tenant_id"),
        ("Ticket",               "tickets",               "id, titolo, tipo (CM|PM|BD), stato, priorita, durata_stimata_ore, tecnico_id, planned_start/finish, is_continuation, deleted_at, tenant_id"),
        ("Manuale",              "manuali",               "id, nome_file, pagine, metodo_lettura, testo_raw, json_estratto, version, stato, tenant_id"),
        ("AttivitaManutenzione", "attivita_manutenzione", "id, asset_id, frequenza_giorni, durata_ore, trigger_mode (calendar|condition), condition_metric, condition_threshold_hours, tenant_id"),
        ("AssetConditionReading","asset_condition_readings","id, asset_id, metric, value, recorded_at — letture ore/cicli"),
        ("AnalisiGuasto",        "analisi_guasto",        "id, ticket_id, metodo, sintomi, analisi_json, ai_utilizzata, tenant_id"),
        ("DiagnosticSession",    "diagnostic_sessions",   "id, ticket_id, history (JSON conversazione), status (active|closed), root_cause, tenant_id"),
        ("TecnicoAssenza",       "tecnico_assenze",       "id, tecnico_id, data_inizio, data_fine, tipo_assenza, tenant_id"),
        ("TicketAllegato",       "ticket_allegati",       "id, ticket_id, nome_file, percorso, tipo_mime, dimensione_bytes, tenant_id"),
        ("EmailConfig",          "email_configs",         "id, imap_server, email_address, password (Fernet), active, default_asset_id, tenant_id"),
        ("AssetDocumento",       "asset_documenti",       "id, asset_id, tipo (Esploso|Manuale|Schema), file_data (BYTEA), esploso_analisi, tenant_id"),
        ("PianoManutenzione",    "piani_manutenzione",    "id, nome_codificato (PM-YYYY-NNNN), progressivo, stato, asset_id, impianto_id, sito_id, tenant_id"),
        ("SystemLog",            "system_logs",           "id, timestamp, level (INFO|WARNING|ERROR), module, message, extra_info, tenant_id"),
        ("GeneratedPlan",        "generated_plans",       "id, status (draft|confirmed|deauthorized), plan_number, plan_json, confirmed_by, deauthorized_by, scadenza, tenant_id"),
        ("PlannerFeedback",      "planner_feedback",      "id, ticket_id, actual_duration_hours, execution_outcome, user_rating (1-5), duration_delta_hours — learning loop"),
        ("FailureMode",          "failure_modes",         "id, asset_type, failure_mode, severity, occurrence, detectability, RPN, peso_appreso, is_global, tenant_id"),
        ("FailureAnalysis",      "failure_analyses",      "id, ticket_id, failure_mode_id, probability_score, rpn_weighted, selected, tenant_id"),
        ("DiagnosticLearning",   "diagnostic_learning",   "id, ticket_id, diagnosed_failure_mode_id, real_cause, action_taken, success, tenant_id"),
        ("Procedura",            "procedure",             "id, asset_id, titolo, tipo (ispezione|sostituzione|loto|emergenza), passi (JSON), revisione, tenant_id"),
        ("NotaAsset",            "note_asset",            "id, asset_id, testo, autore, created_at (upsert), tenant_id"),
        ("CheckPrimoLivello",    "check_primo_livello",   "id, asset_id, public_token (UUID), token_active, token_expires_at, voci (JSON checklist), tenant_id"),
        ("Attestato",            "attestati",             "id, tecnico_id, tipo_corso, ente_certificatore, data_conseguimento, data_scadenza, tenant_id"),
        ("piano_asset_association","piano_asset_association","piano_id + asset_id — M2M piani multi-asset"),
    ],
    col_widths=[1.8, 1.8, 2.9]
)

divider(doc)

# ════════════════════════════════════════════════════════════════════════════
# 16. DEPLOY E INFRASTRUTTURA
# ════════════════════════════════════════════════════════════════════════════
heading(doc, "16. Deploy e Infrastruttura")

ascii_diagram(doc, [
    "  GitHub Repository (branch: main)",
    "       │",
    "       ├────────────────────────────────┐",
    "       │                               │",
    "       ▼                               ▼",
    "  Vercel (Frontend)              Render (Backend)",
    "  ┌──────────────────────┐      ┌────────────────────────────┐",
    "  │ Next.js 16 build     │      │ Python 3.12 + FastAPI       │",
    "  │ Edge CDN globale     │      │ Uvicorn ASGI               │",
    "  │ HTTPS auto (TLS 1.3) │      │ Container web service      │",
    "  │ Deploy automatico    │      │ Porta 10000 (Render)       │",
    "  │ su ogni push main    │      │ Startup: Alembic + init_db │",
    "  └──────────┬───────────┘      └─────────────┬──────────────┘",
    "             │                                │",
    "             │  HTTPS API calls               │  TCP/SSL 5432",
    "             │  Cookie JWT                    ▼",
    "             └──────────────────►   Render PostgreSQL",
    "                                    ┌────────────────────┐",
    "                                    │ Managed DB         │",
    "                                    │ Backup automatico  │",
    "                                    │ SSL enforced       │",
    "                                    └────────────────────┘",
    "                                         │",
    "                                    Supabase Storage",
    "                                    ┌────────────────────┐",
    "                                    │ Bucket privato     │",
    "                                    │ Allegati e PDF     │",
    "                                    │ Signed URL access  │",
    "                                    └────────────────────┘",
    "",
    "  + Desktop App (Tauri 2)",
    "  ┌───────────────────────────────────────────────────────────┐",
    "  │  Rust wrapper + WebView2 (Win) / WebKit (Mac/Linux)       │",
    "  │  Stessa UI Next.js bundled localmente                     │",
    "  │  Auto-update firmato via endpoint /desktop_update         │",
    "  └───────────────────────────────────────────────────────────┘",
], title="Figura 5 — Infrastruttura di deploy MaintAI v3.3.1")

styled_table(doc,
    headers=["Componente", "Provider", "URL produzione", "Note"],
    rows=[
        ("Frontend Web",    "Vercel",    "https://maintai.vercel.app",       "Deploy auto su push main"),
        ("Backend API",     "Render",    "https://maintai-v3.onrender.com",  "Free tier: cold start 30-60s"),
        ("Database",        "Render DB", "Internal network Render",          "PostgreSQL managed, backup daily"),
        ("File Storage",    "Supabase",  "Progetto Supabase (env SUPABASE_URL)", "Bucket privato, signed URL"),
        ("Desktop App",     "GitHub Releases","—",                           "Tauri build Win/Mac/Linux"),
    ],
    col_widths=[1.3, 1.2, 2.4, 1.6]
)

callout(doc, "⚠  ATTENZIONE — Render free tier: il backend può impiegare 30-60 secondi per il "
        "cold start dopo inattività. Il frontend gestisce questa latenza con timeout 120s per "
        "endpoint AI e 30s per gli altri. In produzione considerare upgrade a piano paid Render.")

divider(doc)

# ════════════════════════════════════════════════════════════════════════════
# 17. VARIABILI D'AMBIENTE
# ════════════════════════════════════════════════════════════════════════════
heading(doc, "17. Variabili d'Ambiente")

heading(doc, "17.1 Backend — Render Environment Variables", 3)
styled_table(doc,
    headers=["Variabile", "Obbligatoria", "Descrizione"],
    rows=[
        ("DATABASE_URL",         "Sì",     "URL PostgreSQL Render (postgresql://...)"),
        ("SECRET_KEY",           "Sì",     "Chiave firma JWT (HS256) — min 32 char random"),
        ("ENCRYPTION_KEY",       "Sì",     "Chiave Fernet base64 per cifrare password IMAP"),
        ("OPENAI_API_KEY",       "No",     "Chiave OpenAI — senza: solo motore deterministico"),
        ("OPENAI_MODEL",         "No",     "Modello AI (default: gpt-4.1-mini)"),
        ("CORS_ORIGINS",         "Sì",     "URL frontend autorizzati separati da virgola"),
        ("DEMO_DATABASE_URL",    "No",     "URL SQLite demo (sqlite:///demo.db)"),
        ("SUPABASE_URL",         "No",     "URL progetto Supabase (storage allegati)"),
        ("SUPABASE_KEY",         "No",     "Service role key Supabase"),
        ("ENV",                  "No",     "Se 'production': abilita HSTS, Secure cookie, logging ridotto"),
    ],
    col_widths=[2.0, 1.1, 3.4]
)

heading(doc, "17.2 Frontend — Vercel Environment Variables", 3)
styled_table(doc,
    headers=["Variabile", "Valore esempio", "Descrizione"],
    rows=[
        ("NEXT_PUBLIC_API_BASE",       "https://maintai-v3.onrender.com", "URL base del backend FastAPI"),
        ("NEXT_PUBLIC_DEPLOY_VERSION", "3.3.1",                           "Versione deploy (override version.ts)"),
    ],
    col_widths=[2.2, 2.5, 1.8]
)
callout(doc, "⚠  Le variabili NEXT_PUBLIC_* sono incluse nel bundle JavaScript e visibili nel "
        "browser. Non usarle mai per segreti, chiavi API o credenziali.", RED_ALERT)

divider(doc)

# ════════════════════════════════════════════════════════════════════════════
# 18. LOGGING E OSSERVABILITÀ
# ════════════════════════════════════════════════════════════════════════════
heading(doc, "18. Logging e Osservabilità")

body(doc, "MaintAI usa un sistema di logging a doppio canale: Python standard logging (stdout/file) "
     "e log persistenti nel DB (SystemLog), visibili dalla UI admin.")

styled_table(doc,
    headers=["Canale", "Funzioni", "Dove si legge", "Scopo"],
    rows=[
        ("Python logging",  "logger.info/warning/error",           "Stdout Render / file locale", "Debug tecnico, errori runtime, traceback"),
        ("SystemLog DB",    "log_to_db() / db_info() / db_error()","Frontend /admin/logs",        "Audit trail, visibilità admin, allarmi operativi"),
    ],
    col_widths=[1.5, 2.0, 1.8, 1.2]
)

heading(doc, "18.1 Struttura SystemLog", 3)
styled_table(doc,
    headers=["Campo", "Tipo", "Valori / Descrizione"],
    rows=[
        ("id",         "Integer PK", "Auto-increment"),
        ("tenant_id",  "FK",         "Isolamento log per tenant — ogni tenant vede solo i propri"),
        ("level",      "String",     "INFO / WARNING / ERROR / CRITICAL"),
        ("module",     "String",     "Nome modulo/servizio che ha generato il log"),
        ("message",    "Text",       "Descrizione evento in italiano"),
        ("extra_info", "JSON",       "Payload opzionale (parametri, traceback, context)"),
        ("timestamp",  "DateTime",   "Timestamp UTC evento"),
    ],
    col_widths=[1.3, 1.2, 4.0]
)

heading(doc, "18.2 Health Check", 3)
body(doc, "GET /health — endpoint pubblico (no auth), utilizzabile da monitoring esterno:")
bullet(doc, "status: ok / error")
bullet(doc, "version: versione backend (3.3.1)")
bullet(doc, "build_date: 2026-06-02")
bullet(doc, "database: stato connessione DB (ok / error + dettaglio)")
bullet(doc, "uptime: secondi dall'avvio del processo backend")

doc.add_paragraph()
divider(doc)
doc.add_paragraph()

# ── FOOTER BANNER ─────────────────────────────────────────────────────────
footer_table = doc.add_table(rows=1, cols=1)
footer_table.alignment = WD_TABLE_ALIGNMENT.CENTER
fc = footer_table.cell(0, 0)
set_cell_bg(fc, DARK_BLUE)
para = fc.paragraphs[0]
para.alignment = WD_ALIGN_PARAGRAPH.CENTER
para.paragraph_format.space_before = Pt(8)
para.paragraph_format.space_after = Pt(8)
r = para.add_run(
    "MaintAI v3.3.1  |  Documento di Architettura Tecnica  |  Classificazione: Interno  |  Giugno 2026"
)
r.font.size = Pt(9)
r.font.color.rgb = ACCENT_CYAN

# ── SAVE ─────────────────────────────────────────────────────────────────────
out_path = "/home/user/MAINTAI_MVP_DEMO/MaintAI_Architettura_Tecnica_v3.3.1.docx"
doc.save(out_path)
print(f"✓ Documento salvato: {out_path}")
